from models.llm_client import ask_openai_sync
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document




def convert_text_to_pdf(raw_text, model="gpt-oss-120b"):
    prompt = f"""
    You are a teacher preparing educational content to be turned into a printable PDF. Take the following unformatted text and reformat it using Markdown so it’s clean and readable in a document.

    Use:
    - `#` for main headings
    - `##` for subheadings
    - Bullet points for lists
    - Numbered lists when needed
    - Bold text for key terms or definitions
    - Keep paragraphs clean and separated

    Text:
    \"\"\"{raw_text}\"\"\"
    """
    return ask_openai_sync(prompt=prompt, model=model)


def generate_docx(formatted_text, title=None, class_grade=None, subject=None):
    doc = Document()

    # Add title
    doc.add_heading(title, level=1)

    # Add metadata
    if class_grade or subject:
        meta_line = ""
        if class_grade:
            meta_line += f"Class: {class_grade}  "
        if subject:
            meta_line += f"Subject: {subject}"
        doc.add_paragraph(meta_line)

    doc.add_paragraph("")  # Spacer

    # Add content (split on double newlines)
    for para in formatted_text.strip().split("\n\n"):
        doc.add_paragraph(para.strip())

    # Convert to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()




# PDF generation from Markdown-like formatted text
def generate_pdf(formatted_text, title=None, class_grade=None, subject=None):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    flowables = []

    title_style = styles["Title"]
    subtitle_style = styles["Heading2"]
    normal_style = styles["Normal"]

    # Add metadata
    if title:
        flowables.append(Paragraph(title, title_style))
        flowables.append(Spacer(1, 12))
    
    if class_grade or subject:
        metadata = ""
        if class_grade:
            metadata += f"<b>Class:</b> {class_grade}<br/>"
        if subject:
            metadata += f"<b>Subject:</b> {subject}"
        flowables.append(Paragraph(metadata, subtitle_style))
        flowables.append(Spacer(1, 12))

    # Add formatted text
    for para in formatted_text.split("\n\n"):
        flowables.append(Paragraph(para.strip(), normal_style))
        flowables.append(Spacer(1, 12))

    doc.build(flowables)
    pdf = buffer.getvalue()
    buffer.close()
    return pdf

