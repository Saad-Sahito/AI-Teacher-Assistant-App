from models.llm_client import ask_openai_sync
import os

from docx import Document
from io import BytesIO
import json
import re


def generate_quiz_from_text(text, num_questions, quiz_type="Mixed", class_grade=None, subject=None):
    prompt = f"""
You are a quiz generator AI. Generate {num_questions} questions in {quiz_type} question answer format, from the following study material, for a {subject} {class_grade} class. 
Output questions and answers key in the end.

Study Material:
{text}
"""

    response = ask_openai_sync(prompt)
    #print(response)
    #temp = extract_quiz_json(response)
    #print (temp)
    return response
    

def extract_quiz_json(content_str):
    try:
        content = content_str.strip()

        if content.startswith('['):
            return json.loads(content)

        json_match = re.search(r'\[\s*{.*}\s*\]', content, re.DOTALL)
        if json_match:
            json_str = json_match.group(0)
            return json.loads(json_str)

        raise ValueError("No valid JSON array found in input.")

    except Exception as e:
        return [{
            "question": "Error parsing quiz.",
            "options": [],
            "answer": f"{type(e).__name__}: {str(e)}"
        }]



def generate_quiz_docx(quiz_data, title="Generated Quiz"):
    doc = Document()
    doc.add_heading(title, level=1)

    for idx, q in enumerate(quiz_data, 1):
        doc.add_paragraph(f"Q{idx}: {q['question']}", style='List Number')
        for opt in q['options']:
            doc.add_paragraph(opt, style='List Bullet')
        doc.add_paragraph(f"✅ Answer: {q['answer']}\n")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

